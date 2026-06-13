
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)

# ── Helper: set paragraph spacing ─────────────────────────────────────────────
def set_spacing(para, before=6, after=6):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)

# ── Cover / Title block ───────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("GlobalGive – Transparent Donation Platform")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Aim and Objectives – Detailed Elaboration")
r2.font.size = Pt(14)
r2.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
set_spacing(p2, before=4, after=12)

meta = [
    ("Document Type", "Aim and Objectives – Detailed Elaboration"),
    ("Prepared",      "April 2026"),
]
for k, v in meta:
    pm = doc.add_paragraph()
    pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = pm.add_run(f"{k}: ")
    rm.bold = True
    rm.font.size = Pt(11)
    pm.add_run(v).font.size = Pt(11)
    set_spacing(pm, before=2, after=2)

doc.add_page_break()

# ── Read markdown ─────────────────────────────────────────────────────────────
with open("GlobalGive_Objectives_Elaborated.md", "r") as f:
    lines = f.readlines()

def strip_md(text):
    """Remove inline markdown bold/italic markers."""
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*',     r'\1', text)
    return text.strip()

def add_formatted(para, raw_line):
    """Add runs to a paragraph, applying bold for **text** spans."""
    parts = re.split(r'(\*\*.*?\*\*)', raw_line)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)

i = 0
while i < len(lines):
    line = lines[i].rstrip('\n')

    # Skip raw triple-backtick code blocks (tables inside them etc.) – render as plain
    if line.startswith('```'):
        i += 1
        code_lines = []
        while i < len(lines) and not lines[i].startswith('```'):
            code_lines.append(lines[i].rstrip('\n'))
            i += 1
        # Add code block as styled paragraph
        for cl in code_lines:
            cp = doc.add_paragraph(style='No Spacing')
            run = cp.add_run(cl)
            run.font.name  = 'Courier New'
            run.font.size  = Pt(8.5)
            run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
            cp.paragraph_format.left_indent = Inches(0.4)
            set_spacing(cp, before=0, after=0)
        i += 1
        continue

    # H1
    if line.startswith('# ') and not line.startswith('##'):
        p = doc.add_heading(line[2:].strip(), level=1)
        p.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        set_spacing(p, before=18, after=8)

    # H2
    elif line.startswith('## '):
        p = doc.add_heading(line[3:].strip(), level=2)
        p.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        set_spacing(p, before=14, after=6)

    # H3
    elif line.startswith('### '):
        p = doc.add_heading(line[4:].strip(), level=3)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x1D, 0x4E, 0x89)
        set_spacing(p, before=10, after=4)

    # Horizontal rule
    elif line.strip() in ('---', '___', '***'):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1A56DB')
        pBdr.append(bottom)
        pPr.append(pBdr)
        set_spacing(p, before=4, after=4)

    # Markdown table rows — render as plain paragraph
    elif line.startswith('|'):
        cells = [c.strip() for c in line.strip('|').split('|')]
        # Skip separator rows like |---|---|
        if all(re.match(r'^-+$', c) for c in cells if c):
            i += 1
            continue
        p = doc.add_paragraph(style='No Spacing')
        p.paragraph_format.left_indent = Inches(0.2)
        text = '   '.join(cells)
        add_formatted(p, text)
        p.runs[0].font.size = Pt(10) if p.runs else None
        set_spacing(p, before=1, after=1)

    # Bullet / list items
    elif line.startswith('- ') or line.startswith('* '):
        content = line[2:].strip()
        p = doc.add_paragraph(style='List Bullet')
        add_formatted(p, strip_md(content) if '**' not in content else content.replace('**',''))
        p.runs[0].font.size = Pt(11)
        set_spacing(p, before=2, after=2)

    # Numbered list
    elif re.match(r'^\d+\.\s', line):
        content = re.sub(r'^\d+\.\s', '', line).strip()
        p = doc.add_paragraph(style='List Number')
        add_formatted(p, content)
        p.runs[0].font.size = Pt(11)
        set_spacing(p, before=2, after=2)

    # Page break marker
    elif 'doc.add_page_break' in line or line.strip() == '---page---':
        doc.add_page_break()

    # Blank line
    elif line.strip() == '':
        pass

    # Normal paragraph
    else:
        p = doc.add_paragraph()
        add_formatted(p, line)
        for run in p.runs:
            run.font.size = Pt(11)
        p.paragraph_format.first_line_indent = Pt(0)
        set_spacing(p, before=3, after=3)

    i += 1

# ── Save ──────────────────────────────────────────────────────────────────────
output = "GlobalGive_Objectives_Elaborated.docx"
doc.save(output)
print(f"✅ DOCX saved: {output}")
